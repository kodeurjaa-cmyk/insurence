from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from io import BytesIO
import markdown2 # Need to add this to requirements if used, otherwise simple text
import os

class Exporter:
    def __init__(self):
        pass

    def export_to_docx(self, policy_text: str, filename: str) -> str:
        doc = Document()
        doc.add_heading('Insurance Policy', 0)
        
        # Simple breakdown of markdown-like text
        for line in policy_text.split('\n'):
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            else:
                doc.add_paragraph(line)
        
        filepath = os.path.join('exports', f"{filename}.docx")
        os.makedirs('exports', exist_ok=True)
        doc.save(filepath)
        return filepath

    def export_to_pdf(self, policy_text: str, filename: str) -> str:
        filepath = os.path.join('exports', f"{filename}.pdf")
        os.makedirs('exports', exist_ok=True)
        
        c = canvas.Canvas(filepath, pagesize=letter)
        width, height = letter
        
        text_object = c.beginText(40, height - 40)
        text_object.setFont("Helvetica", 10)
        
        lines = policy_text.split('\n')
        for line in lines:
            if text_object.getY() < 40:
                c.drawText(text_object)
                c.showPage()
                text_object = c.beginText(40, height - 40)
                text_object.setFont("Helvetica", 10)
            
            text_object.textLine(line)
        
        c.drawText(text_object)
        c.save()
        return filepath
